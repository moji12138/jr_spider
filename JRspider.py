#!/usr/bin/env python3
# _*_ coding:utf-8 _*_
'''
version: 0.1
author: moji
Python version: Python3.4+
功能：将今日头条的文章和悟空问答的页面生成docx格式的文档
头条文章 √
悟空问答 √
头条图集 ×
付费文章 ×
'''
import re
import os
import hashlib

import requests
from requests.exceptions import RequestException
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


class JinRi():
    '''
    工厂类，用来根据url信息生成实例
    '''

    __obj = None
    __init_flag = True

    @staticmethod
    def get(url):
        '''
        静态方法，根据url参数返回Wukong()或者Toutiao()实例
        '''
        # 判断今日头条url中的关键字
        url_rule = r"toutiao|pstatp|wukong"
        if re.search(url_rule, url, re.I):
            header = {
                'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.86 Safari/537.36'
            }
            reqobj = requests.get(url, headers=header)
            surl_rule = r"([^0-9]*\d*/)"
            surl = surl = re.search(surl_rule,reqobj.url).group(1)
            if reqobj.url.count("wukong"):
                return Wukong(surl)
            else:
                return Toutiao(surl)

    def __new__(cls):
        if cls.__obj == None:
            cls.__obj = object.__new__(cls)
        return cls.__obj

    def __init__(self):
        if JinRi.__init_flag:
            JinRi.__init_flag = False


class JR2docx():
    '''
    今日头条类 Wukong类和Toutiao的父类
    '''

    def __init__(self, url):
        '''
        初始化成员属性
        链接地址 url
        请求头 header
        文档标题 title
        正文内容 data
        '''
        self.url = url
        self.header = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.86 Safari/537.36'
        }
        self.title = "未获得标题..."
        self.data = self.get_data()

    def __str__(self):
        return "%s[%s]" % (self.title,self.url)

    __repr__ = __str__

    @property
    def get_requests(self):
        '''
        请求网页数据方法
        请求网页返回reqests对象
        '''
        try:
            req = requests.get(self.url, headers=self.header)
            req.encoding = r"utf-8"
            if req.status_code == 200:
                return req.text
            return None
        except RequestException:
            print(u"请求页面出错")
            return None

    def __download_img(self, url):
        '''
        下载图片方法
        请求图片url并保存图片
        返回图片路径
        '''
        img = requests.get(url, headers=self.header)
        if img.status_code == 200:
            # 图片文件路径 当前路径下的img文件夹 文件名为文件的md5 格式为jpg
            filepath = os.path.join(os.getcwd(), "img", hashlib.md5(
                img.content).hexdigest() + ".jpg")
            with open(filepath, 'wb') as fp:
                fp.write(img.content)
                fp.close
        else:
            return None
        return filepath

    def get_data(self):
        '''
        获得有效数据
        '''
        raise NotImplementedError

    def get_docx(self):
        '''
        生成文档的方法
        返回文档名
        '''
        docName = self.title + ".docx"  # 文档名
        document = Document()
        document.styles['Normal'].font.name = u'宋体'  # 设置正文为宋体
        document.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋体')
        document.add_heading(self.title, level=0)  # 插入0级标题
        try:
            for insert in self.data:
                # 如果包含关键字"pstatp"则此内容为图片链接 调用下载图片方法
                if insert.count("pstatp"):
                    # 判断是否以http开头，否则加上http补全链接
                    if insert.count("http") == 0:
                        insert = "http:" + insert
                    imgPath = self.__download_img(insert)
                    # 如果图片路径不为空 插入图片
                    if imgPath:
                        document.add_picture(imgPath, width=Inches(4))
                else:
                    # 插入文字
                    document.add_paragraph(insert)
            document.add_paragraph("本文链接地址为：\n(%s)" % self.url)  # 插入本文url地址
            document.add_page_break()  # 插入结束符
            document.save(r"document/%s" % docName)  # 保存文件
        except TypeError:
            return None
        return docName


class Toutiao(JR2docx):
    '''
    头条文章类
    继承今日头条类
    '''
    def get_data(self):
        '''
        返回有效数据的字典
        '''
        textRule = re.compile(r"(src&#x3D;&quot;|&gt;)(.*?)(&quot;|&lt;)")   # 正文文本规则
        titleRule = re.compile(r"title: '(.*?)',")                           # 标题规则
        webText = []    # 初始化正文列表
        try:
            searchTitle = titleRule.search(self.get_requests)
            self.title = searchTitle.group(1)
            content = textRule.finditer(self.get_requests)
            # 如果内容不为空 添加到正文列表
            for text in content:
                if text.group(2) is not "":
                    webText.append(text.group(2))
            
            return webText
        except TypeError:
            print("爬取数据失败")
            return None


class Wukong(JR2docx):
    '''
    悟空问答类
    继承今日头条类
    '''
    def get_data(self):
        '''
        返回有效数据的字典
        '''
        soup = BeautifulSoup(self.get_requests, 'lxml')
        try:
            dataSoup = soup.select(
                r"#main-index-question-list > div.question.question-single > div > div.all-answers > div.answers > div.answer-items > div > div.answer-text.h_1200 > div.answer-text-full.rich-text")[0]
            title = soup.select(
                r"#main-index-question-list > div.question.question-single > div > div > h1 > a")[0].text
            self.title = title
            webText = []
            # 遍历数据列表 如果不为空 插入数据列表
            for p in dataSoup:
                if p.text:
                    webText.append(p.text)
                if p.get("src"):
                    webText.append(p.get("src"))
                if p.img:
                    webText.append(p.img.get("src"))
            
            return webText
        except TypeError:
            print("爬取数据失败...")
            return None


if __name__ == "__main__":
    url = "https://www.toutiao.com/a6682926082330460684/"
    jrobj = JinRi.get(url)
    print(jrobj.get_docx())
    # url = "https://m.zjurl.cn/answer/6676996556949815565/?app=news_article&amp;app_id=13&amp;share_ansid=6676996556949815565&amp;wxshare_count=1&amp;tt_from=weixin&amp;utm_source=weixin&amp;utm_medium=toutiao_android&amp;utm_campaign=client_share"
    # jrobj = JinRi.get(url)
    # print(jrobj.get_docx())
    # url = "https://m.zjurl.cn/answer/6676992333466042632/?app=news_article&amp;app_id=13&amp;share_ansid=6676992333466042632&amp;wxshare_count=1&amp;tt_from=weixin&amp;utm_source=weixin&amp;utm_medium=toutiao_android&amp;utm_campaign=client_share"
    # jrobj = JinRi.get(url)
    # print(jrobj.get_docx())
